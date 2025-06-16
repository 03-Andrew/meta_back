from docx import Document
from openpyxl import Workbook
import os
import random
from faker import Faker

fake = Faker()
os.makedirs("test-files/docx", exist_ok=True)
os.makedirs("test-files/xlsx", exist_ok=True)

def generate_docx_file(path, word_count):
    doc = Document()
    doc.add_heading(fake.catch_phrase(), 0)

    # Add body paragraphs
    for _ in range(word_count // 100):
        doc.add_paragraph(fake.text(max_nb_chars=1000))

    # Add metadata
    core_props = doc.core_properties
    core_props.author = fake.name()
    core_props.title = fake.bs()
    core_props.subject = fake.job()
    core_props.keywords = ", ".join(fake.words(nb=5))
    core_props.comments = fake.sentence()
    doc.save(path)

def generate_xlsx_file(path, rows, cols):
    wb = Workbook()
    ws = wb.active

    for r in range(1, rows + 1):
        for c in range(1, cols + 1):
            ws.cell(row=r, column=c, value=fake.word())

    props = wb.properties
    props.creator = fake.name()
    props.title = fake.catch_phrase()
    props.subject = fake.company()
    props.keywords = ", ".join(fake.words(nb=5))
    props.description = fake.sentence()
    wb.save(path)

# Generate small, medium, large .docx and .xlsx
sizes = {
    'small': (300, 50, 5),
    'medium': (3000, 300, 10),
    'large': (15000, 1000, 20)
}

for size, (words, rows, cols) in sizes.items():
    for i in range(1, 4):
        generate_docx_file(f"test-files/docx/{size}_{i}.docx", word_count=words)
        generate_xlsx_file(f"test-files/xlsx/{size}_{i}.xlsx", rows=rows, cols=cols)

